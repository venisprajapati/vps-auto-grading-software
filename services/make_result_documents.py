from datetime import date
from docx import Document
from docx.shared import Mm


def ResultDocument(exam_name_, exam_date_, section_, results, omr_s_, students, total_marks, section_1_n='', section_2_n='', section_3_n=''):

    document = Document()

    # A4 Paper
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(12.7)
    section.right_margin = Mm(12.7)
    section.top_margin = Mm(12.7)
    section.bottom_margin = Mm(12.7)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)

    # header
    header = section.header

    header_para = header.paragraphs[0]

    header_para.text = f"Venis Prajapati's Auto grading Software\t\t{date.today()}"

    # footer
    footer = section.footer

    footer_para = footer.paragraphs[0]

    footer_para.text = f"{str(exam_name_)} {str(exam_date_)}\t\tgithub: venisprajapati/vps-auto-grading-software"

    # title
    document.add_heading("Venis Prajapati's Auto grading Software", 0)

    # Exam name
    document.add_heading(
        f'Exam: {str(exam_name_)}\t\t\tDate: {str(exam_date_)}\t\t\tMarks: {total_marks}', 1)

    # heading 1
    document.add_heading("Result Analysis Report by Marks", 2)

    # rank, id, name, correct, not attempted, incorrect, marks
    table = document.add_table(rows=1, cols=7)
    table.style = 'Light Grid Accent 4'
    row = table.rows[0]
    row.cells[0].text = 'Rank'
    row.cells[1].text = 'ID'
    row.cells[2].text = 'Name'
    row.cells[3].text = 'Correct'
    row.cells[4].text = 'N/A'
    row.cells[5].text = 'Incorrect'
    row.cells[6].text = 'Marks Obt'

    for result in results:
        cells = table.add_row().cells
        cells[0].text = str(result['rank'])
        cells[1].text = str(result['id'])
        cells[2].text = students[result['id']] if students.get(
            result['id']) is not None else str(result['id'])
        cells[3].text = str(result['correct_count'])
        cells[4].text = str(result['not_attempted_count'])
        cells[5].text = str(result['incorrect_count'])
        cells[6].text = str(result['obtained_total_marks'])

    if (section_):

        # heading 2
        document.add_heading("Result Analysis Report by Sections", 2)

        # rank, id, name, section-1, section-2, section-3, marks
        analysis_table = document.add_table(rows=1, cols=7)
        analysis_table.style = 'Light Grid Accent 4'
        row = analysis_table.rows[0]
        row.cells[0].text = 'Rank'
        row.cells[1].text = 'ID'
        row.cells[2].text = 'Name'
        row.cells[3].text = section_1_n
        row.cells[4].text = section_2_n
        row.cells[5].text = section_3_n
        row.cells[6].text = 'Marks Obt'

        for result in results:
            cells = analysis_table.add_row().cells
            cells[0].text = str(result['rank'])
            cells[1].text = str(result['id'])
            cells[2].text = students[result['id']] if students.get(
                result['id']) is not None else str(result['id'])
            cells[3].text = str(result['section_1_marks']
                                if result.get('section_1') else 'N/A')
            cells[4].text = str(result['section_2_marks']
                                if result.get('section_2') else 'N/A')
            cells[5].text = str(result['section_3_marks']
                                if result.get('section_3') else 'N/A')
            cells[6].text = str(result['obtained_total_marks'])

    # make result sheets
    document.add_page_break()

    # heading 3
    document.add_heading("Results for Individual Students", 2)

    document.add_paragraph(
        "________________________________________________________________________________________________________________________________")

    for result in results:

        # Starting of paragraph
        head_paragraph = document.add_paragraph()
        head_paragraph.add_run('Exam: ').bold = True
        head_paragraph.add_run(str(exam_name_))
        head_paragraph.add_run('\t\tDate: ').bold = True
        head_paragraph.add_run(str(exam_date_))
        head_paragraph.add_run('\t\tMax marks: ').bold = True
        head_paragraph.add_run(str(total_marks))

        head_paragraph.add_run('\t\tID: ').bold = True
        head_paragraph.add_run(str(result['id']))

        main_paragraph = document.add_paragraph()
        main_paragraph.add_run('Rank: ').bold = True
        main_paragraph.add_run(str(result['rank']))
        main_paragraph.add_run('\tName: ').bold = True
        main_paragraph.add_run(students[result['id']] if students.get(
            result['id']) is not None else str(result['id']))
        main_paragraph.add_run('\tCorrect: ').bold = True
        main_paragraph.add_run(str(result['correct_count']))
        main_paragraph.add_run('\tN/A: ').bold = True
        main_paragraph.add_run(str(result['not_attempted_count']))
        main_paragraph.add_run('\t\tIncorrect: ').bold = True
        main_paragraph.add_run(str(result['incorrect_count']))
        main_paragraph.add_run('\tMarks Obt: ').bold = True
        main_paragraph.add_run(str(result['obtained_total_marks']))

        if (section_):
            main_paragraph.add_run(f'\n{section_1_n}: ').bold = True
            main_paragraph.add_run(
                str(result['section_1_marks'] if result.get('section_1') else 'N/A'))
            main_paragraph.add_run(f'\t\t\t{section_2_n}: ').bold = True
            main_paragraph.add_run(
                str(result['section_2_marks'] if result.get('section_2') else 'N/A'))
            main_paragraph.add_run(f'\t\t\t{section_3_n}: ').bold = True
            main_paragraph.add_run(
                str(result['section_3_marks'] if result.get('section_3') else 'N/A'))

        main_paragraph.add_run('\nCorrect: ').bold = True
        main_paragraph.add_run(str(result['correct']))
        main_paragraph.add_run(' N/A: ').bold = True
        main_paragraph.add_run(str(result['not_attempted']))
        main_paragraph.add_run(' Incorrect: ').bold = True
        main_paragraph.add_run(str(result['incorrect']))

        document.add_paragraph(
            "________________________________________________________________________________________________________________________________")

    # heading 3
    document.add_heading("Checked OMRs Of Individual Students", 2)

    document.add_paragraph(
        "________________________________________________________________________________________________________________________________")

    for omr in omr_s_:
        omr_par_ = document.add_paragraph()
        omr_par_.add_run('ID: ').bold = True
        omr_par_.add_run(str(omr['id']))
        omr_par_.add_run(' ')

        for i_mcq_ in range(1, len(omr)):
            omr_par_.add_run(str(i_mcq_)).bold = True
            omr_par_.add_run(str(f'{omr[i_mcq_]},'))

        document.add_paragraph(
            "________________________________________________________________________________________________________________________________")

    document.save(
        f'./result/{exam_name_.replace(" ", "_").replace("-", "_")}_{exam_date_.replace("-", "_")}.docx')
